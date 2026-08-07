"""
Document Processing & Extraction Module

Supports parsing PDF, DOCX, TXT, Markdown, CSV, and HTML/Web pages.
Preserves page numbers, metadata, and cleans extracted text.
"""

import io
import re
import uuid
from typing import Any, Dict, List, Optional

from app.core.logging import get_logger
from app.schemas.research import Document

logger = get_logger(__name__)

# Graceful optional imports for parsing libraries
try:
    import fitz  # PyMuPDF for PDF
except ImportError:
    fitz = None

try:
    import docx  # python-docx for Word docs
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None


class DocumentLoader:
    """Multi-format Document Processing Service."""

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text content."""
        if not text:
            return ""
        # Replace non-breaking spaces and normalize whitespace
        text = text.replace("\xa0", " ").replace("\r\n", "\n")
        # Collapse multiple blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    async def load_from_text(
        self,
        text: str,
        title: str = "Raw Document",
        file_type: str = "txt",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Process raw string text (TXT or Markdown)."""
        logger.info("Loading document from raw text", title=title, file_type=file_type)
        cleaned = self.clean_text(text)
        doc_id = str(uuid.uuid4())
        return Document(
            document_id=doc_id,
            title=title,
            content=cleaned,
            file_type=file_type,
            source_path_or_url=title,
            page_count=1,
            metadata=metadata or {},
        )

    async def load_from_html(
        self,
        html_content: str,
        title: str = "Web Page",
        url: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Extract clean text from HTML content using BeautifulSoup."""
        logger.info("Loading document from HTML", title=title, url=url)
        if BeautifulSoup:
            soup = BeautifulSoup(html_content, "html.parser")
            # Remove scripts and styling elements
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
        else:
            # Fallback regex HTML tag stripping
            text = re.sub(r"<[^>]+>", " ", html_content)

        cleaned = self.clean_text(text)
        return Document(
            document_id=str(uuid.uuid4()),
            title=title,
            content=cleaned,
            file_type="html",
            source_path_or_url=url or title,
            page_count=1,
            metadata=metadata or {},
        )

    async def load_from_pdf_bytes(
        self,
        pdf_bytes: bytes,
        title: str = "Document.pdf",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Extract text and page metadata from PDF using PyMuPDF (fitz)."""
        logger.info("Loading document from PDF bytes", title=title)
        pages_text: List[str] = []
        page_count = 1

        if fitz:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page_count = len(doc)
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                txt = page.get_text()
                pages_text.append(f"[Page {page_num + 1}]\n{txt}")
            combined_text = "\n\n".join(pages_text)
        else:
            logger.warning("PyMuPDF (fitz) not installed, extracting printable ASCII fallback")
            combined_text = pdf_bytes.decode("utf-8", errors="ignore")

        cleaned = self.clean_text(combined_text)
        return Document(
            document_id=str(uuid.uuid4()),
            title=title,
            content=cleaned,
            file_type="pdf",
            source_path_or_url=title,
            page_count=page_count,
            metadata=metadata or {},
        )

    async def load_from_docx_bytes(
        self,
        docx_bytes: bytes,
        title: str = "Document.docx",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Extract text from DOCX bytes using python-docx."""
        logger.info("Loading document from DOCX bytes", title=title)
        if docx:
            doc_file = io.BytesIO(docx_bytes)
            doc_obj = docx.Document(doc_file)
            paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
            combined = "\n\n".join(paragraphs)
        else:
            logger.warning("python-docx not installed, using fallback string decoding")
            combined = docx_bytes.decode("utf-8", errors="ignore")

        cleaned = self.clean_text(combined)
        return Document(
            document_id=str(uuid.uuid4()),
            title=title,
            content=cleaned,
            file_type="docx",
            source_path_or_url=title,
            page_count=1,
            metadata=metadata or {},
        )

    async def load_from_csv_bytes(
        self,
        csv_bytes: bytes,
        title: str = "Data.csv",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Document:
        """Extract text from CSV bytes using pandas or built-in string reader."""
        logger.info("Loading document from CSV bytes", title=title)
        if pd:
            df = pd.read_csv(io.BytesIO(csv_bytes))
            combined = df.to_string()
        else:
            combined = csv_bytes.decode("utf-8", errors="ignore")

        cleaned = self.clean_text(combined)
        return Document(
            document_id=str(uuid.uuid4()),
            title=title,
            content=cleaned,
            file_type="csv",
            source_path_or_url=title,
            page_count=1,
            metadata=metadata or {},
        )
